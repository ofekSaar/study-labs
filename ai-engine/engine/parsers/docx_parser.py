"""
Word Document Parser — Extracts text, tables, and embedded images from DOCX files.

Uses python-docx for structured XML parsing.
No OCR needed — Word documents store text as structured data.
"""

import io
import logging
from typing import List

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):

    @staticmethod
    def supported_extensions() -> List[str]:
        return ['.docx']

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        logger.info(f"DocxParser: Processing '{filename}'...")
        text_parts = []
        images = []

        try:
            doc = Document(io.BytesIO(file_bytes))

            # --- Paragraph extraction ---
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve heading structure
                    if para.style and para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '').strip()
                        try:
                            level_num = int(level)
                            text = "#" * level_num + " " + text
                        except ValueError:
                            text = "## " + text
                    text_parts.append(text)

            # --- Table extraction ---
            for table in doc.tables:
                table_rows = []
                for row_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append("| " + " | ".join(row_cells) + " |")

                if table_rows:
                    # Insert Markdown header separator after first row
                    num_cols = len(table.rows[0].cells)
                    header_sep = "| " + " | ".join(["---"] * num_cols) + " |"
                    table_rows.insert(1, header_sep)
                    text_parts.append("\n".join(table_rows))

            # --- Image extraction ---
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_blob = rel.target_part.blob
                        if image_blob:
                            images.append(image_blob)
                            logger.debug(f"  Extracted embedded image from '{filename}'")
            except Exception as e:
                logger.warning(f"  Image extraction failed for '{filename}': {e}")

        except Exception as e:
            logger.error(f"DocxParser Error for '{filename}': {e}")
            return ParseResult(text="", images=[], metadata={"error": str(e)})

        full_text = "\n\n".join(text_parts)
        logger.info(f"DocxParser: Extracted {len(full_text)} chars, {len(images)} images from '{filename}'")

        return ParseResult(
            text=full_text,
            images=images,
            metadata={"paragraph_count": len(text_parts), "image_count": len(images)}
        )
